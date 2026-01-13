#!/usr/bin/env python3
"""Convert CODE_GUIDE_FOR_BEGINNERS.md to Word document"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([__import__("sys").executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read markdown file
with open('CODE_GUIDE_FOR_BEGINNERS.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create Word document
doc = Document()

# Add title
title = doc.add_heading('Web Development Code Guide - NPA Level 5', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Simple Explanations for Beginners')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.bold = True
subtitle_format.font.size = Pt(14)

doc.add_paragraph()  # Space

# Convert markdown to Word (simplified)
lines = content.split('\n')
for line in lines:
    line = line.strip()
    
    if not line:
        continue
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('| '):
        # Skip table markers for now
        pass
    elif line.startswith('```'):
        # Skip code fence markers
        pass
    elif line.startswith('**') or '**' in line:
        # Add as bold paragraph
        p = doc.add_paragraph()
        p.add_run(line).bold = True
    else:
        doc.add_paragraph(line)

# Save document
doc.save('CODE_GUIDE_FOR_BEGINNERS.docx')
print("✓ Word document created: CODE_GUIDE_FOR_BEGINNERS.docx")
